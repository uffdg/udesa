"""
Kit de post-procesado para cualquier .docx generado con pandoc +
design-system/components/docx-reference.docx (ver docx-guion.md en esta
misma carpeta para el flujo completo).

Por qué existe: docx-reference.docx define el look de tabla (header oscuro,
filas alternadas, bordes hairline BORDER) vía w:tblStylePr (firstRow /
band1Horz) en el estilo de tabla "Table" — es la forma "correcta" de
OOXML, y Microsoft Word la respeta. Pero Pages (verificado acá,
2026-08-04, Pages de macOS) no aplica el shading/color de esas
conditional formats de tabla al importar un .docx ajeno: respeta bordes y
peso de fuente del header, pero ignora el fill oscuro y el color de fuente
blanco del header, y no pinta las bandas alternadas. En vez de dejar ese
riesgo (verse bien en Word, verse plano en Pages/otros lectores con
soporte OOXML más débil), este script aplica el mismo resultado visual
directo en cada celda — shading y color de fuente sobre la celda misma,
no sobre un estilo de tabla — que ningún lector de .docx puede ignorar.

Uso:
    python3 design-system/components/docx_kit.py archivo.docx
    (in-place; solo toca tablas, no toca texto/headings/blockquotes, que sí
    se resuelven bien vía docx-reference.docx solo)
"""
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

INK = "222222"
BG_LIGHT = "F8F8F8"
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    # sacar shading previo (el que pudiera venir del estilo de tabla) para
    # no dejar dos <w:shd> compitiendo
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_font(cell, *, bold: bool = None, color: RGBColor = None) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if bold is not None:
                run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color


def style_tables(doc_path: str) -> int:
    """Aplica header oscuro + bandas alternadas directo sobre las celdas de
    cada tabla del documento. Devuelve la cantidad de tablas procesadas."""
    doc = Document(doc_path)
    n = 0
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        # fila 0 = header (pandoc siempre marca w:tblHeader="on" en la
        # primera fila de cualquier tabla pipe/grid de markdown)
        for cell in rows[0].cells:
            _set_cell_shading(cell, INK)
            _set_cell_font(cell, bold=True, color=WHITE_RGB)
        # bandas alternadas en el resto: impar (1ra fila de cuerpo) = BG_LIGHT
        for idx, row in enumerate(rows[1:]):
            if idx % 2 == 0:
                for cell in row.cells:
                    _set_cell_shading(cell, BG_LIGHT)
        n += 1
    doc.save(doc_path)
    return n


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("uso: python3 docx_kit.py archivo.docx", file=sys.stderr)
        sys.exit(1)
    path = sys.argv[1]
    if not Path(path).exists():
        print(f"no existe: {path}", file=sys.stderr)
        sys.exit(1)
    count = style_tables(path)
    print(f"OK — {count} tabla(s) formateadas en {path}")
